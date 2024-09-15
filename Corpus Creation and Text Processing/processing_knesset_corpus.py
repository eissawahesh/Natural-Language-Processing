import re
import time
import os
from docx import Document
import pandas as pd

class Protocol:
    def __init__(self, file_name, knesset_number, protocol_type, file_path):
        self.file_name=file_name
        self.knesset_number = knesset_number
        self.protocol_type = protocol_type
        self.file_path = file_path
        self.list = []


def check_hebrewchar_and_digit(char):
    return '\u0590' <= char <= '\u05FF' or char.isdigit()

def tokenization(sentence):
    marks = [".", "?", "!","(",")"]
    tokens = sentence.split()
    new_tokens = []

    for token in tokens:
        if check_hebrewchar_and_digit(token[0]) and check_hebrewchar_and_digit(token[-1]):
            new_tokens.append(token)
        elif not check_hebrewchar_and_digit(token[0]):
            if token[0] in marks:
                new_tokens.append(token[0])
                new_tokens.append(token[1:])
            else:
                new_tokens.append(token[1:])
        elif not check_hebrewchar_and_digit(token[-1]):
            if token[-1] in marks:
                new_tokens.append(token[:-1])
                new_tokens.append(token[-1])
            else:
                new_tokens.append(token[:-1])
    new_sentence = ' '.join(new_tokens)
    return new_sentence


def split_into_sentences(text):
    sentences = re.split(r'(?<=[.!?])', text)
    new_sentences=[]
    for sentence in sentences:
        new_sentence=sentence.strip()
        if new_sentence:
            new_sentences.append(new_sentence)
    return new_sentences

def clean_ending_sentence(speaker_speeches_file):
    text = speaker_speeches_file[-1][1]
    index1 = text.find("הישיבה ננעלה")
    if index1 != -1:
        speaker_speeches_file[-1][1] = text[:index1]
    index2 = text.find("הטקס ננעל בשעה")
    if index2 != -1:
        speaker_speeches_file[-1][1] = text[:index2]

def speakers_and_speeches_ptv(file_path):
    document = Document(file_path)
    speakers_speeches = []
    current_speaker = None
    counter = -1

    for para in document.paragraphs:
        if para.text.strip().endswith(":>")or para.text.strip().endswith(":") or para.text.strip().endswith(">>"):
            if para.style.name == "דובר" or para.style.name== "יור" or para.style.name=="דובר-המשך" or para.style.name== "קריאות" or para.style.name=="אורח":
                counter += 1
                current_speaker = para.text
                speakers_speeches.append([current_speaker, ""])

            else:
                for run in para.runs:
                    if run.underline and not run.bold:
                        counter += 1
                        current_speaker = para.text
                        speakers_speeches.append([current_speaker, ""])
                        break
        if current_speaker  and para.text!=current_speaker :
            speakers_speeches[counter][1] += para.text + "\n"

    clean_ending_sentence(speakers_speeches)

    return speakers_speeches



def speakers_and_speeches_ptm(file_path):
    document = Document(file_path)
    speakers_speeches = []
    current_speaker = None
    counter = -1

    for para in document.paragraphs:
        if para.text.strip().endswith(":"):
            if para.style.name == "דובר" or para.style.name== "יור" or para.style.name=="דובר-המשך" or para.style.name=="קריאות" :
                counter += 1
                current_speaker = para.text
                speakers_speeches.append([current_speaker, ""])

            else:
                for run in para.runs:
                    if run.underline and run.bold:
                        counter += 1
                        current_speaker = para.text
                        speakers_speeches.append([current_speaker, ""])
                        break
        if current_speaker and para.style.name == "Normal" and para.text!=current_speaker:
            speakers_speeches[counter][1] += para.text + "\n"
    clean_ending_sentence(speakers_speeches)

    return speakers_speeches

def cleaning_sentences(speakers_and_speeches_list):
    new_list=[]

    for speaker in speakers_and_speeches_list:
        if not speaker[1].strip():
            continue
        if(check_english_char(speaker[1])):
            continue
        if("- -" in speaker[1]):
            continue
        if ("– –" in speaker[1]):
            continue
        if(speaker[1].startswith("הצבעה")):
            continue
        new_list.append(speaker)
    return new_list

def check_english_char(sentence):
    for char in sentence:
        if 'A' <= char <= 'Z' or 'a' <= char <= 'z':
            return True
    return False


def cleaning_names(speakers_list):
    for speaker in speakers_list:
        if ("קריאה") in speaker[0] or  ("קריאות") in speaker[0] or not speaker[0].strip():
            speakers_list.remove(speaker)
    characters = ["(", ")", '"']
    special_cases = ["<", ">", ":", "יור ", "אורח ", "דובר_המשך", "דובר","סגן שר במשרד ראש הממשלה","המשנה לראש הממשלה",
        "ראש הממשלה",'יו"ר ועדת הכנסת',"נשיא הפרלמנט האירופי",'יו"ר ועדת העבודה, הרווחה והבריאות',"שר התחבורה והבטיחות בדרכים ",
         "שר התשתיות הלאומיות, האנרגיה והמים ", "שר העבודה, הרווחה והשירותים החברתיים" ,"שר הרווחה והשירותים החברתיים"
       , "השר לאיכות הסביבה","השר לאיכות הסביבה" , "השר לביטחון הפנים " ,"השר לקליטת העלייה","השר לשיתוף פעולה אזורי ","שר התשתיות הלאומיות, האנרגיה והמים"
       , "השר לביטחון פנים" , "היו”ר ", "השר לנושאים אסטרטגיים ולענייני מודיעין ",'היו"ר',"פרופ'"
         ,  "שר החקלאות ופיתוח הכפר", "השר להגנת הסביבה", "שר התשתיות הלאומיות ", 'יו"ר הכנסת'
        , "השר לאזרחים ותיקים ", "השר "]
    for s in range(len(speakers_list)):
        for case in special_cases:
            speakers_list[s][0] = speakers_list[s][0].replace(case,"")

        if ("תשובת") in speakers_list[s][0]:
            remains = speakers_list[s][0].split()
            remains = [word for word in remains if word not in ["תשובת"]]
            speakers_list[s][0]= ' '.join(remains)

        if ("סגן") in speakers_list[s][0] or  ("סגנית") in speakers_list[s][0]:
            remains = speakers_list[s][0].split()
            remains = [word for word in remains if word not in ["סגן","סגנית"]]
            speakers_list[s][0]= ' '.join(remains)
        if ("מזכיר הכנסת") in speakers_list[s][0] or  ("מזכירת הכנסת") in speakers_list[s][0]:
            remains = speakers_list[s][0].split()
            remains = [word for word in remains if word not in ["מזכיר הכנסת","מזכירת הכנסת"]]
            speakers_list[s][0] = ' '.join(remains[2:])
        if ("שר") in speakers_list[s][0] or ("שרת") in speakers_list[s][0]:
            remains = speakers_list[s][0].split()
            remains = [word for word in remains if word not in ["שר", "שרת"]]


            if len(remains)>1:
                if remains[0].startswith("ה"):
                    if remains[0].endswith(","):
                        if remains[1].startswith("ה") and remains[2].startswith("ו"):
                            speakers_list[s][0] = ' '.join((remains[3:]))

                    else:
                        if len(remains)>2 and remains[1].startswith("וה") :
                            speakers_list[s][0] = ' '.join((remains[2:]))
                        else:
                            speakers_list[s][0] = ' '.join((remains[1:]))




        first_char = speakers_list[s][0].find('(')
        last_char = speakers_list[s][0].find(')')
        if first_char != -1 and last_char != -1:
            speakername=speakers_list[s][0][:first_char]+speakers_list[s][0][last_char+1:]
            speakers_list[s][0]=speakername.strip()

        words=speakers_list[s][0].split()
        cleaned=[]

        for word in words:
          if not any(char in word for char in characters):
              cleaned.append(word)

        cleaned_name=' '.join(cleaned)
        speakers_list[s][0]=cleaned_name
start_time=time.time()
files_list = []

def process_corpus(input_path,output_path):
    directory_path = input_path
    for filename in os.listdir(directory_path):
        full_path = os.path.join(directory_path, filename)
        part1, part2, part3 = filename.split("_")
        part3 = full_path
        if part2 == "ptv":
            part2 = "committee"
        else:
            part2 = "plenary"

        files_list.append(Protocol(file_name=filename, knesset_number=part1, protocol_type=part2, file_path=part3))

    for i in range(100):
        file = files_list[i]
        if (file.protocol_type == "committee"):
            speakers_speeches = speakers_and_speeches_ptv(file.file_path)
        elif (file.protocol_type == "plenary"):
            speakers_speeches = speakers_and_speeches_ptm(file.file_path)

        cleaning_names(speakers_speeches)
        speech_for_speaker = []
        for speaker_speech in speakers_speeches:
            sentences = split_into_sentences(speaker_speech[1])

            for sentence in sentences:
                speech_for_speaker.append([speaker_speech[0], sentence])

        speech_for_speaker = cleaning_sentences(speech_for_speaker)
        for sp in speech_for_speaker:
            file.list.append([sp[0], tokenization(sp[1])])

    data = []

    for i in range(100):
        for speaker_speech in files_list[i].list:
            tokens = speaker_speech[1].split()
            if len(tokens) >= 4:
                data.append({
                    "protocol_name": files_list[i].file_name,
                    "kneseet_number": files_list[i].knesset_number,
                    "protocol_type": files_list[i].protocol_type,
                    "speaker_name": speaker_speech[0],
                    "sentence_text": speaker_speech[1]
                })

    df = pd.DataFrame(data)

    df.to_csv(output_path, encoding='utf-8-sig', index=False)

import argparse

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_path")
    parser.add_argument("output_path")
    args = parser.parse_args()
    process_corpus(args.input_path,args.output_path)
main()